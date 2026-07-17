import os
import time
import uuid
import base64
import pandas as pd
from typing import List, Dict, Any, Optional
from pypdf import PdfReader
from docx import Document as DocxDocument
from PIL import Image
import chromadb
from chromadb.utils import embedding_functions
from google import genai
from google.genai import types
from groq import Groq

# Define local directories
UPLOAD_DIR = os.path.abspath("./uploaded_files")
CHROMA_DIR = os.path.abspath("./chroma_db")

os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(CHROMA_DIR, exist_ok=True)

class DocumentProcessor:
    @staticmethod
    def extract_text_from_txt(file_path: str) -> str:
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        except Exception as e:
            return f"Error reading TXT file: {str(e)}"

    @staticmethod
    def extract_text_from_pdf(file_path: str, api_key: Optional[str] = None) -> str:
        try:
            reader = PdfReader(file_path)
            text = ""
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text += f"\n--- Page {i+1} ---\n" + page_text
            text = text.strip()
            
            # Fallback to PyMuPDF and Groq Vision OCR if the text layer is empty or extremely short
            if len(text.replace("\n", "").strip()) < 50 and api_key:
                import fitz
                doc = fitz.open(file_path)
                ocr_text = []
                # Capped at first 10 pages for safety
                max_pages = min(len(doc), 10)
                for idx in range(max_pages):
                    page = doc.load_page(idx)
                    pix = page.get_pixmap(dpi=150)
                    img_data = pix.tobytes("png")
                    
                    base64_image = base64.b64encode(img_data).decode('utf-8')
                    try:
                        from groq import Groq
                        client = Groq(api_key=api_key)
                        response = client.chat.completions.create(
                            model="meta-llama/llama-4-scout-17b-16e-instruct",
                            messages=[
                                {
                                    "role": "user",
                                    "content": [
                                        {
                                            "type": "text",
                                            "text": "Transcribe all text from this page image exactly and completely."
                                        },
                                        {
                                            "type": "image_url",
                                            "image_url": {
                                                "url": f"data:image/png;base64,{base64_image}"
                                            }
                                        }
                                    ]
                                }
                            ],
                            max_tokens=1000
                        )
                        page_ocr = response.choices[0].message.content
                        if page_ocr:
                            ocr_text.append(f"\n--- Page {idx+1} (OCR) ---\n" + page_ocr)
                    except Exception as e:
                        ocr_text.append(f"\n--- Page {idx+1} (OCR Failed) ---\nError: {e}")
                
                if ocr_text:
                    text = "\n".join(ocr_text).strip()
            return text
        except Exception as e:
            return f"Error reading PDF file: {str(e)}"

    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        try:
            doc = DocxDocument(file_path)
            text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text.append(" | ".join(row_text))
            return "\n".join(text)
        except Exception as e:
            return f"Error reading DOCX file: {str(e)}"

    @staticmethod
    def extract_text_from_excel(file_path: str) -> str:
        try:
            xls = pd.ExcelFile(file_path)
            text = []
            for sheet_name in xls.sheet_names:
                df = pd.read_excel(xls, sheet_name=sheet_name)
                text.append(f"Sheet Name: {sheet_name}")
                text.append(f"Columns: {', '.join(df.columns.astype(str))}")
                
                # Format each row compactly, up to 500 rows to ensure fast indexing
                limit = 500
                row_strings = []
                for idx, row in df.head(limit).iterrows():
                    row_parts = []
                    for col, val in row.items():
                        if pd.notna(val):
                            row_parts.append(f"{col}: {val}")
                    if row_parts:
                        row_strings.append(f"Row {idx+1}: " + " | ".join(row_parts))
                
                text.extend(row_strings)
                if len(df) > limit:
                    text.append(f"... (Truncated {len(df) - limit} rows for indexing performance)")
            return "\n".join(text)
        except Exception as e:
            return f"Error reading Excel file: {str(e)}"

    @staticmethod
    def extract_text_from_image(file_path: str, api_key: Optional[str] = None) -> str:
        # Standard metadata description
        metadata_desc = ""
        try:
            with Image.open(file_path) as img:
                metadata_desc = f"Image Format: {img.format} | Size: {img.size[0]}x{img.size[1]} | Mode: {img.mode}"
        except Exception as e:
            metadata_desc = f"Error reading Image file metadata: {str(e)}"
            
        if not api_key:
            return metadata_desc

        # If we have an API key, we run Groq Vision OCR
        try:
            with open(file_path, "rb") as image_file:
                base64_image = base64.b64encode(image_file.read()).decode('utf-8')
                
            client = Groq(api_key=api_key)
            chat_completion = client.chat.completions.create(
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": "Perform high-accuracy OCR on this image. Extract all readable text, names, numbers, tables, and marksheet details. Output ONLY the extracted text and structured tables, no introduction or chat commentary."},
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/jpeg;base64,{base64_image}",
                                },
                            },
                        ],
                    }
                ],
                model="meta-llama/llama-4-scout-17b-16e-instruct",
            )
            extracted_text = chat_completion.choices[0].message.content
            # Return metadata + extracted text
            return f"{metadata_desc}\n\n=== Extracted Image Text (OCR) ===\n{extracted_text}"
        except Exception as e:
            return f"{metadata_desc}\n\n(Groq Vision OCR failed: {str(e)})"

    @classmethod
    def process_file(cls, file_path: str, file_type: str, api_key: Optional[str] = None) -> Dict[str, Any]:
        """Extracts content representation and metadata from a file."""
        content = ""
        metadata = {
            "filename": os.path.basename(file_path),
            "type": file_type,
            "timestamp": float(time.time()),
        }
        
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".txt":
            content = cls.extract_text_from_txt(file_path)
        elif ext == ".pdf":
            content = cls.extract_text_from_pdf(file_path, api_key)
        elif ext in [".docx"]:
            content = cls.extract_text_from_docx(file_path)
        elif ext in [".xls", ".xlsx"]:
            content = cls.extract_text_from_excel(file_path)
        elif ext in [".png", ".jpg", ".jpeg", ".gif", ".webp"]:
            content = cls.extract_text_from_image(file_path, api_key)
        else:
            content = "Unsupported file format content."
            
        return {"content": content, "metadata": metadata}

    @staticmethod
    def chunk_text(text: str, chunk_size: int = 1000, chunk_overlap: int = 200) -> List[str]:
        if not text:
            return []
        
        chunks = []
        start = 0
        while start < len(text):
            end = start + chunk_size
            chunks.append(text[start:end])
            start += chunk_size - chunk_overlap
            
        return chunks

class LocalVectorStore:
    def __init__(self):
        # We try to initialize chromadb with default persistent settings
        try:
            self.client = chromadb.PersistentClient(path=CHROMA_DIR)
            self.collection = self.client.get_or_create_collection(
                name="doc_agent_collection",
                metadata={"hnsw:space": "cosine"}
            )
            self.use_fallback = False
        except Exception as e:
            print(f"Error initializing ChromaDB: {e}. Falling back to simple dictionary store.")
            self.use_fallback = True
            self.fallback_db = {} # simple memory store

    def add_document(self, file_id: str, content: str, doc_metadata: Dict[str, Any]):
        # Chunk text
        if doc_metadata["type"] in ["XLS", "XLSX"]:
            chunks = []
            lines = content.split("\n")
            current_chunk = []
            current_len = 0
            for line in lines:
                current_chunk.append(line)
                current_len += len(line)
                if len(current_chunk) >= 15 or current_len >= 1000:
                    chunks.append("\n".join(current_chunk))
                    current_chunk = []
                    current_len = 0
            if current_chunk:
                chunks.append("\n".join(current_chunk))
        else:
            chunks = DocumentProcessor.chunk_text(content)

        if not chunks:
            # If no content, at least add a metadata stub
            chunks = [f"File metadata: {doc_metadata['filename']} is a {doc_metadata['type']} uploaded at {doc_metadata['timestamp']}."]

        if self.use_fallback:
            self.fallback_db[file_id] = {
                "chunks": chunks,
                "metadata": doc_metadata
            }
            return

        # Add to Chroma
        ids = [f"{file_id}_chunk_{i}" for i in range(len(chunks))]
        metadatas = [doc_metadata.copy() for _ in range(len(chunks))]
        # Add chunk index to metadata
        for i, meta in enumerate(metadatas):
            meta["chunk_index"] = i
            
        self.collection.add(
            documents=chunks,
            metadatas=metadatas,
            ids=ids
        )

    def search(self, query: str, n_results: int = 5) -> List[Dict[str, Any]]:
        if self.use_fallback or not query.strip():
            # Basic keyword matching fallback search
            results = []
            q_words = set(query.lower().split())
            if not q_words:
                return []
                
            for file_id, doc in self.fallback_db.items():
                for idx, chunk in enumerate(doc["chunks"]):
                    score = sum(1 for w in q_words if w in chunk.lower())
                    if score > 0:
                        results.append({
                            "document": chunk,
                            "metadata": {**doc["metadata"], "chunk_index": idx},
                            "score": float(score)
                        })
            # sort descending by score
            results = sorted(results, key=lambda x: x["score"], reverse=True)
            return results[:n_results]

        try:
            res = self.collection.query(
                query_texts=[query],
                n_results=n_results
            )
            
            # format results
            formatted = []
            if res and res["documents"] and len(res["documents"]) > 0:
                docs = res["documents"][0]
                metas = res["metadatas"][0]
                distances = res["distances"][0] if res["distances"] else [0.0] * len(docs)
                for i in range(len(docs)):
                    # convert distance to a similarity score (cosine distance to similarity: 1 - dist)
                    sim = 1.0 - distances[i]
                    formatted.append({
                        "document": docs[i],
                        "metadata": metas[i],
                        "score": sim
                    })
            return formatted
        except Exception as e:
            print(f"Query error: {e}")
            return []

    def delete_document(self, filename: str):
        if self.use_fallback:
            # remove from fallback_db dictionary
            keys_to_remove = [k for k, doc in self.fallback_db.items() if doc["metadata"]["filename"] == filename]
            for k in keys_to_remove:
                del self.fallback_db[k]
            return
            
        try:
            self.collection.delete(where={"filename": filename})
        except Exception as e:
            print(f"Delete document error: {e}")

    def reset_db(self):
        if self.use_fallback:
            self.fallback_db.clear()
            return
        try:
            self.client.delete_collection("doc_agent_collection")
            self.collection = self.client.get_or_create_collection(
                name="doc_agent_collection",
                metadata={"hnsw:space": "cosine"}
            )
        except Exception as e:
            print(f"Reset database failed: {e}")

# Groq Answer Synthesizer
def generate_groq_answer(api_key: str, query: str, context_chunks: List[Dict[str, Any]]) -> str:
    """Generates an answer using the Groq SDK if API key is provided."""
    if not api_key:
        return "Please configure your Groq API key in the sidebar to generate AI answers."
        
    try:
        # Build prompt
        context_str = ""
        for i, chunk in enumerate(context_chunks):
            context_str += f"\n[{i+1}] Source: {chunk['metadata']['filename']} (Chunk {chunk['metadata'].get('chunk_index', 0)}):\n{chunk['document']}\n"
            
        prompt = f"""You are an intelligent AI Document Assistant. Synthesize a clear, accurate, and comprehensive answer to the User's Query using the provided retrieved context.
If the context does not contain the answer, specify that you cannot find it in the uploaded documents, but provide a best-effort answer based on the available information.
Always mention which documents/files the answer is sourced from.

User's Query: {query}

Retrieved Context:
{context_str}

Answer:"""
        
        # Initialize client with user provided API key
        client = Groq(api_key=api_key)
        chat_completion = client.chat.completions.create(
            messages=[
                {
                    "role": "user",
                    "content": prompt,
                }
            ],
            model="llama-3.3-70b-versatile",
        )
        return chat_completion.choices[0].message.content
    except Exception as e:
        return f"Error generating answer with Groq: {str(e)}"
